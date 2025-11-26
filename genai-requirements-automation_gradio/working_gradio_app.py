#!/usr/bin/env python3
import gradio as gr
import sys
import os
sys.path.append('backend')

from pdf_processor import PDFProcessor
from openai_service import OpenAIService
from document_generator import DocumentGenerator
from azure_devops_service import AzureDevOpsService
from dotenv import load_dotenv

load_dotenv()

# Global variables
generated_documents = None

def extract_pdf_text(pdf_file):
    """Extract text from uploaded PDF"""
    try:
        if pdf_file is None:
            return ""
        
        with open(pdf_file.name, 'rb') as f:
            pdf_text = PDFProcessor.extract_text_from_pdf(f)
        return PDFProcessor.clean_text(pdf_text)
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"

def generate_documents(pdf_text, gen_epic, gen_story, gen_usecase, gen_tech, gen_data, gen_brd, gen_fsd, gen_test):
    """Generate selected documents"""
    global generated_documents
    
    try:
        if not pdf_text.strip():
            return "No content to process", "", "", "", "", "", "", "", ""
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key or api_key == "your_openai_api_key_here":
            return "Configure OpenAI API key", "", "", "", "", "", "", "", ""
        
        openai_service = OpenAIService(api_key)
        documents = openai_service.generate_documents(pdf_text)
        generated_documents = documents
        
        # Format outputs based on selections
        outputs = [""] * 8
        
        if gen_epic and documents.epics:
            outputs[0] = "\n".join([f"# {epic.title}\n{epic.description}\n**Business Value:** {epic.business_value}\n**Acceptance Criteria:** {epic.acceptance_criteria}\n" for epic in documents.epics])
        
        if gen_story and documents.stories:
            outputs[1] = "\n".join([f"# {story.title}\n{story.description}\n**Points:** {story.story_points}\n**Acceptance Criteria:** {story.acceptance_criteria}\n" for story in documents.stories])
        
        if gen_usecase and documents.use_cases:
            outputs[2] = "\n".join([f"# {uc.title}\n**Actor:** {uc.actor}\n**Preconditions:** {uc.preconditions}\n**Main Flow:** {', '.join(uc.main_flow)}\n**Postconditions:** {uc.postconditions}\n" for uc in documents.use_cases])
        
        if gen_tech:
            tech = documents.technical_design
            outputs[3] = f"# {tech.title}\n**Overview:** {tech.overview}\n**Architecture:** {tech.architecture}\n**Components:** {', '.join(tech.components)}\n**Technologies:** {', '.join(tech.technologies)}\n**Data Flow:** {tech.data_flow}\n**Security:** {tech.security_considerations}"
        
        if gen_data:
            data = documents.data_model
            entities = "\n".join([f"- {entity['name']}: {', '.join(entity['attributes'])}" for entity in data.entities])
            relationships = "\n".join([f"- {rel['from']} -> {rel['to']} ({rel['type']})" for rel in data.relationships])
            outputs[4] = f"# {data.title}\n**Entities:**\n{entities}\n**Relationships:**\n{relationships}\n**Constraints:** {', '.join(data.constraints)}"
        
        if gen_brd and documents.business_requirement_doc:
            brd = documents.business_requirement_doc
            outputs[5] = f"# {brd.title}\n**Executive Summary:** {brd.executive_summary}\n**Business Objectives:** {', '.join(brd.business_objectives)}\n**Functional Requirements:** {', '.join(brd.functional_requirements)}\n**Non-Functional Requirements:** {', '.join(brd.non_functional_requirements)}"
        
        if gen_fsd and documents.functional_spec_doc:
            fsd = documents.functional_spec_doc
            functions = "\n".join([f"- {func['name']}: {func['description']}" for func in fsd.system_functions])
            outputs[6] = f"# {fsd.title}\n**Overview:** {fsd.overview}\n**System Functions:**\n{functions}\n**Business Rules:** {', '.join(fsd.business_rules)}"
        
        if gen_test and documents.test_cases:
            outputs[7] = "\n".join([f"# {tc.test_name} ({tc.test_id})\n**Description:** {tc.description}\n**Preconditions:** {tc.preconditions}\n**Steps:** {', '.join(tc.test_steps)}\n**Expected Result:** {tc.expected_result}\n**Priority:** {tc.priority}\n" for tc in documents.test_cases])
        
        return "Documents generated successfully!", *outputs
        
    except Exception as e:
        return f"Error: {str(e)}", "", "", "", "", "", "", "", ""

def create_download_file(epic_text, story_text, usecase_text, tech_text, data_text, brd_text, fsd_text, test_text, dl_epic, dl_story, dl_usecase, dl_tech, dl_data, dl_brd, dl_fsd, dl_test, format_type):
    """Create download file with selected content"""
    try:
        content = ""
        
        if dl_epic and epic_text:
            content += f"# EPICS\n\n{epic_text}\n\n"
        if dl_story and story_text:
            content += f"# USER STORIES\n\n{story_text}\n\n"
        if dl_usecase and usecase_text:
            content += f"# USE CASES\n\n{usecase_text}\n\n"
        if dl_tech and tech_text:
            content += f"# TECHNICAL DESIGN\n\n{tech_text}\n\n"
        if dl_data and data_text:
            content += f"# DATA MODEL\n\n{data_text}\n\n"
        if dl_brd and brd_text:
            content += f"# BUSINESS REQUIREMENTS\n\n{brd_text}\n\n"
        if dl_fsd and fsd_text:
            content += f"# FUNCTIONAL SPECIFICATION\n\n{fsd_text}\n\n"
        if dl_test and test_text:
            content += f"# TEST CASES\n\n{test_text}\n\n"
        
        if not content:
            return None
        
        if format_type == "PDF":
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                
                filename = "generated_documents.pdf"
                doc = SimpleDocTemplate(filename, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                for line in content.split('\n'):
                    if line.startswith('#'):
                        story.append(Paragraph(line.replace('#', ''), styles['Heading1']))
                    elif line.strip():
                        story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 12))
                
                doc.build(story)
                return filename
            except ImportError:
                filename = "generated_documents.txt"
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                return filename
        
        elif format_type == "Word":
            try:
                from docx import Document
                
                filename = "generated_documents.docx"
                doc = Document()
                
                for line in content.split('\n'):
                    if line.startswith('#'):
                        doc.add_heading(line.replace('#', '').strip(), level=1)
                    elif line.strip():
                        doc.add_paragraph(line)
                
                doc.save(filename)
                return filename
            except ImportError:
                filename = "generated_documents.md"
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                return filename
        
    except Exception as e:
        return None

def push_to_ado():
    """Push epics and stories to Azure DevOps"""
    global generated_documents
    
    try:
        if not generated_documents:
            return "No documents generated yet"
        
        if not all([os.getenv("AZURE_DEVOPS_ORG_URL"), os.getenv("AZURE_DEVOPS_PAT"), os.getenv("AZURE_DEVOPS_PROJECT")]):
            return "Configure Azure DevOps credentials in .env file"
        
        azure_service = AzureDevOpsService(
            os.getenv("AZURE_DEVOPS_ORG_URL"),
            os.getenv("AZURE_DEVOPS_PAT"),
            os.getenv("AZURE_DEVOPS_PROJECT")
        )
        
        results = azure_service.push_all_items(generated_documents.epics, generated_documents.stories, generated_documents.use_cases)
        
        summary = f"Pushed {len(results['epics'])} epics, {len(results['stories'])} stories, {len(results['use_cases'])} use cases"
        if results['errors']:
            summary += f"\nErrors: {len(results['errors'])}\n"
            for error in results['errors']:
                summary += f"- {error}\n"
        
        return summary
        
    except Exception as e:
        return f"Error pushing to ADO: {str(e)}"

def simple_chat(message, history):
    """Simple chat without complex formatting"""
    global generated_documents
    
    if not generated_documents:
        return "Please generate documents first before asking questions."
    
    # Simple response based on generated content
    if "epic" in message.lower():
        if generated_documents.epics:
            return f"Found {len(generated_documents.epics)} epics: " + ", ".join([epic.title for epic in generated_documents.epics])
        else:
            return "No epics found in generated documents."
    elif "story" in message.lower():
        if generated_documents.stories:
            return f"Found {len(generated_documents.stories)} stories: " + ", ".join([story.title for story in generated_documents.stories])
        else:
            return "No stories found in generated documents."
    else:
        return "I can help you with questions about epics and stories. Try asking 'What epics were generated?' or 'What stories were created?'"

# Create Gradio interface
with gr.Blocks(title="GenAI Document Generator") as app:
    gr.Markdown("# GenAI Document Generator")
    
    with gr.Row():
        with gr.Column():
            # File upload
            pdf_file = gr.File(label="Upload PDF", file_types=[".pdf"])
            pdf_text = gr.Textbox(label="Extracted PDF Content", lines=8, placeholder="PDF content will appear here...")
            
            # Generation options
            gr.Markdown("### Select Documents to Generate:")
            with gr.Row():
                gen_epic = gr.Checkbox(label="Epic", value=True)
                gen_story = gr.Checkbox(label="Story", value=True)
                gen_usecase = gr.Checkbox(label="Use Case", value=True)
            with gr.Row():
                gen_tech = gr.Checkbox(label="Technical Design", value=True)
                gen_data = gr.Checkbox(label="Data Model", value=True)
            with gr.Row():
                gen_brd = gr.Checkbox(label="Business Requirement Doc", value=True)
                gen_fsd = gr.Checkbox(label="Functional Spec Doc", value=True)
                gen_test = gr.Checkbox(label="Test Cases", value=True)
            
            generate_btn = gr.Button("Generate Documents", variant="primary")
            gen_status = gr.Textbox(label="Generation Status", interactive=False)
        
        with gr.Column():
            # Simple chat
            gr.Markdown("### Document Assistant")
            chat_interface = gr.ChatInterface(
                simple_chat,
                title="Ask about your documents",
                description="Ask questions about the generated documents"
            )
    
    # Generated content areas
    with gr.Row():
        with gr.Column():
            epic_output = gr.Textbox(label="Epics", lines=4)
            story_output = gr.Textbox(label="Stories", lines=4)
            usecase_output = gr.Textbox(label="Use Cases", lines=4)
            tech_output = gr.Textbox(label="Technical Design", lines=4)
        with gr.Column():
            data_output = gr.Textbox(label="Data Model", lines=4)
            brd_output = gr.Textbox(label="Business Requirements", lines=4)
            fsd_output = gr.Textbox(label="Functional Specification", lines=4)
            test_output = gr.Textbox(label="Test Cases", lines=4)
    
    with gr.Row():
        with gr.Column():
            # Download options
            gr.Markdown("### Download Options:")
            with gr.Row():
                dl_epic = gr.Checkbox(label="Epic", value=True)
                dl_story = gr.Checkbox(label="Story", value=True)
                dl_usecase = gr.Checkbox(label="Use Case", value=True)
                dl_tech = gr.Checkbox(label="Technical Design", value=True)
            with gr.Row():
                dl_data = gr.Checkbox(label="Data Model", value=True)
                dl_brd = gr.Checkbox(label="Business Requirements", value=True)
                dl_fsd = gr.Checkbox(label="Functional Spec", value=True)
                dl_test = gr.Checkbox(label="Test Cases", value=True)
            
            format_type = gr.Radio(["PDF", "Word"], label="Format", value="PDF")
            download_btn = gr.Button("Download Documents")
            download_file = gr.File(label="Download", interactive=False)
        
        with gr.Column():
            # ADO integration
            gr.Markdown("### Azure DevOps Integration:")
            ado_btn = gr.Button("Push to ADO", variant="secondary")
            ado_status = gr.Textbox(label="ADO Status", interactive=False)
    
    # Event handlers
    pdf_file.change(fn=extract_pdf_text, inputs=[pdf_file], outputs=[pdf_text])
    
    generate_btn.click(
        fn=generate_documents,
        inputs=[pdf_text, gen_epic, gen_story, gen_usecase, gen_tech, gen_data, gen_brd, gen_fsd, gen_test],
        outputs=[gen_status, epic_output, story_output, usecase_output, tech_output, data_output, brd_output, fsd_output, test_output]
    )
    
    download_btn.click(
        fn=create_download_file,
        inputs=[epic_output, story_output, usecase_output, tech_output, data_output, brd_output, fsd_output, test_output,
                dl_epic, dl_story, dl_usecase, dl_tech, dl_data, dl_brd, dl_fsd, dl_test, format_type],
        outputs=[download_file]
    )
    
    ado_btn.click(fn=push_to_ado, outputs=[ado_status])

if __name__ == "__main__":
    app.launch(share=False)