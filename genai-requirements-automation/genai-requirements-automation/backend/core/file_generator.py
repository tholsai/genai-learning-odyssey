"""File generator for creating DOCX and PDF documents."""
import os
import uuid
from typing import Dict, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from core.config import settings


class FileGenerator:
    """Generate DOCX and PDF files from content."""
    
    @staticmethod
    def _ensure_directory(directory: str):
        """Ensure directory exists."""
        os.makedirs(directory, exist_ok=True)
    
    @staticmethod
    def generate_docx(
        content: str,
        title: str,
        output_path: str
    ) -> str:
        """Generate a DOCX file from content."""
        FileGenerator._ensure_directory(os.path.dirname(output_path))
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a heading (starts with #)
                if para_text.strip().startswith('#'):
                    level = len(para_text) - len(para_text.lstrip('#'))
                    heading_text = para_text.lstrip('#').strip()
                    doc.add_heading(heading_text, level=min(level, 9))
                else:
                    para = doc.add_paragraph(para_text.strip())
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def generate_pdf(
        content: str,
        title: str,
        output_path: str
    ) -> str:
        """Generate a PDF file from content."""
        FileGenerator._ensure_directory(os.path.dirname(output_path))
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Build PDF content
        story = []
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add content
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a heading
                if para_text.strip().startswith('#'):
                    level = len(para_text) - len(para_text.lstrip('#'))
                    heading_text = para_text.lstrip('#').strip()
                    style_name = f'Heading{min(level, 6)}'
                    story.append(Paragraph(heading_text, styles[style_name]))
                    story.append(Spacer(1, 0.1*inch))
                else:
                    story.append(Paragraph(para_text.strip(), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        return output_path
    
    @staticmethod
    def generate_artifact_file(
        content: str,
        artifact_type: str,
        doc_type: str = "docx"
    ) -> str:
        """Generate a file for a specific artifact type."""
        FileGenerator._ensure_directory(settings.generated_dir)
        
        # Map artifact types to titles
        titles = {
            "epic": "Epic Document",
            "stories": "User Stories",
            "use_cases": "Use Cases",
            "tdd": "Test-Driven Development Test Cases",
            "data_model": "Data Model"
        }
        
        title = titles.get(artifact_type, "Generated Document")
        filename = f"{artifact_type}_{uuid.uuid4().hex[:8]}.{doc_type}"
        output_path = os.path.join(settings.generated_dir, filename)
        
        if doc_type.lower() == "docx":
            return FileGenerator.generate_docx(content, title, output_path)
        elif doc_type.lower() == "pdf":
            return FileGenerator.generate_pdf(content, title, output_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

